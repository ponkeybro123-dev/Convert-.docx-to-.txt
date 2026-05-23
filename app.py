from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
import os
from werkzeug.utils import secure_filename
import io

app = Flask(__name__)

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx'}
MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB

# Create uploads folder if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE


def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def convert_docx_to_txt(docx_path):
    """Convert a .docx file to .txt format"""
    try:
        # Load the docx file
        doc = Document(docx_path)
        
        # Extract text from all paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                text_content.append(' | '.join(row_text))
        
        return '\n'.join(text_content)
    except Exception as e:
        raise Exception(f"Error converting file: {str(e)}")


@app.route('/')
def index():
    """Render the home page"""
    return render_template('index.html')


@app.route('/convert', methods=['POST'])
def convert():
    """Handle file upload and conversion"""
    try:
        # Check if file is in request
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400
        
        file = request.files['file']
        
        # Check if file is selected
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
        
        # Check if file is allowed
        if not allowed_file(file.filename):
            return jsonify({'error': 'Only .docx files are allowed'}), 400
        
        # Save the uploaded file
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Convert the file
        text_content = convert_docx_to_txt(filepath)
        
        # Create output filename
        output_filename = filename.rsplit('.', 1)[0] + '.txt'
        
        # Clean up the uploaded file
        os.remove(filepath)
        
        return jsonify({
            'success': True,
            'filename': output_filename,
            'content': text_content
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/download', methods=['POST'])
def download():
    """Download the converted .txt file"""
    try:
        data = request.get_json()
        content = data.get('content', '')
        filename = data.get('filename', 'converted.txt')
        
        # Create a BytesIO object
        output = io.BytesIO()
        output.write(content.encode('utf-8'))
        output.seek(0)
        
        return send_file(
            output,
            mimetype='text/plain',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True, port=5000)
